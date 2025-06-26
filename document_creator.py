from docx import Document
from docx.shared import Pt, Inches
from io import BytesIO
from fpdf import FPDF
import re

def format_bold(text, paragraph):
    
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        run = paragraph.add_run(part.replace('**', '')) if '**' in part else paragraph.add_run(part)
        if '**' in part:
            run.bold = True

def create_resume_docx(resume_text):
    
    document = Document()

    for section in document.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    lines = [line.strip() for line in resume_text.strip().split('\n') if line.strip()]

    if lines:
        title = document.add_heading(lines[0], level=1)
        title.runs[0].font.size = Pt(16)
        lines.pop(0)

    headings = {"Resumo Profissional", "Experiência Profissional", "Formação Acadêmica", "Habilidades"}

    for line in lines:
        if line in headings:
            p = document.add_paragraph()
            p.add_run(line).bold = True
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
        elif line.startswith('•') or line.startswith('-'):
            document.add_paragraph(line, style='List Bullet')
        else:
            
            para = document.add_paragraph()
            if '|' in line and '•' not in line and 'Habilidades' not in line:
                para.paragraph_format.space_before = Pt(8)
            format_bold(line, para)

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer

def create_resume_pdf(resume_text):
    
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(False) 

    def to_latin1(text):
        return text.encode('latin-1', 'replace').decode('latin-1')

    def clean_bold(text):
        return re.sub(r'\*\*(.*?)\*\*', r'\1', text)

    line_height = 5

    def check_page_break():
        
        if pdf.get_y() > pdf.page_break_trigger - (line_height * 2): 
            pdf.add_page()

    def write_with_wrap(text, indent=0):
        
        check_page_break()
        pdf.set_x(pdf.l_margin + indent)
        words = text.split(' ')
        
        for word in words:
            word_to_write = word + ' '
            word_width = pdf.get_string_width(word_to_write)
            
            if pdf.get_x() + word_width > pdf.w - pdf.r_margin:
                pdf.ln(line_height)
                check_page_break()
                pdf.set_x(pdf.l_margin + indent)
            
            pdf.cell(word_width, line_height, word_to_write)
        
        pdf.ln(line_height) 

    lines = [line.strip() for line in resume_text.strip().split('\n') if line.strip()]
    
    if lines:
        name = to_latin1(clean_bold(lines.pop(0)))
        pdf.set_font("Arial", 'B', 16)
        pdf.multi_cell(0, 10, name, align='C')
        pdf.ln(5)

    headings = {"Resumo Profissional", "Experiência Profissional", "Formação Acadêmica", "Habilidades"}
    
    for line in lines:
        check_page_break()
        clean_line = to_latin1(clean_bold(line))

        if clean_line in headings:
            pdf.ln(4)
            pdf.set_font("Arial", 'B', 12)
            pdf.cell(0, 10, clean_line, ln=1, align='L')
        elif clean_line.startswith(('•', '-', '*')):
            pdf.set_font("Arial", '', 11)
            bullet_indent = 5
            text_indent = 10
            
            pdf.set_x(pdf.l_margin + bullet_indent)
            pdf.cell(text_indent - bullet_indent, line_height, '\x95')
            
            text_content = clean_line.lstrip('•-* ').strip()
            write_with_wrap(text_content, indent=text_indent)
        else:
           
            if '|' in clean_line and '•' not in clean_line and 'Habilidades' not in clean_line:
                pdf.ln(3)  
            pdf.set_font("Arial", '', 11)
            write_with_wrap(clean_line)

    return bytes(pdf.output())
